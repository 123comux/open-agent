"""Tests for document loaders (PDF, DOCX, CSV, JSON, HTML, Markdown, Text)."""
from __future__ import annotations

import csv
import json
from pathlib import Path

import pytest

from open_agent.rag.document_loaders import (
    SUPPORTED_EXTENSIONS,
    load_csv,
    load_docx,
    load_file,
    load_html,
    load_json,
    load_markdown,
    load_pdf,
    load_text,
)


def test_supported_extensions() -> None:
    """All expected extensions are supported."""
    expected = {
        ".txt", ".md", ".rst", ".text", ".pdf", ".docx", ".csv", ".json", ".html", ".htm"
    }
    assert expected.issubset(SUPPORTED_EXTENSIONS)


def test_load_text(tmp_path: Path) -> None:
    """Plain text files are loaded verbatim."""
    file_path = tmp_path / "sample.txt"
    file_path.write_text("Hello world", encoding="utf-8")
    doc = load_text(str(file_path))
    assert doc.text == "Hello world"
    assert doc.file_type == "text"
    assert doc.metadata["source"] == str(file_path)


def test_load_markdown_strips_frontmatter(tmp_path: Path) -> None:
    """YAML frontmatter is stripped from Markdown files."""
    file_path = tmp_path / "sample.md"
    file_path.write_text("---\ntitle: Test\n---\n\nBody text", encoding="utf-8")
    doc = load_markdown(str(file_path))
    assert "title: Test" not in doc.text
    assert "Body text" in doc.text
    assert doc.file_type == "markdown"


def test_load_pdf(tmp_path: Path) -> None:
    """PDF loader extracts text with page markers and metadata."""
    try:
        import fitz  # type: ignore[import-untyped]
    except ImportError:  # pragma: no cover
        pytest.skip("PyMuPDF not installed")

    file_path = tmp_path / "sample.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Hello from page one")
    doc.save(str(file_path))
    doc.close()

    loaded = load_pdf(str(file_path))
    assert "Hello from page one" in loaded.text
    assert "--- Page 1 ---" in loaded.text
    assert loaded.metadata["pages"] == 1
    assert loaded.file_type == "pdf"


def test_load_docx(tmp_path: Path) -> None:
    """DOCX loader extracts headings, paragraphs and tables."""
    try:
        import docx
    except ImportError:  # pragma: no cover
        pytest.skip("python-docx not installed")

    file_path = tmp_path / "sample.docx"
    document = docx.Document()
    document.add_heading("Introduction", level=1)
    document.add_paragraph("This is a sample paragraph.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "A"
    table.cell(1, 1).text = "1"
    document.save(str(file_path))

    loaded = load_docx(str(file_path))
    assert "# Introduction" in loaded.text
    assert "This is a sample paragraph." in loaded.text
    assert "Name | Value" in loaded.text
    assert "A | 1" in loaded.text
    assert loaded.metadata["headings"] >= 1
    assert loaded.metadata["tables"] == 1
    assert "Introduction" in loaded.metadata.get("heading_titles", [])
    assert loaded.file_type == "docx"


def test_load_csv(tmp_path: Path) -> None:
    """CSV loader converts rows to a Markdown-like table string."""
    file_path = tmp_path / "sample.csv"
    with open(file_path, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["Name", "Age"])
        writer.writerow(["Alice", "30"])
        writer.writerow(["Bob", "25"])

    loaded = load_csv(str(file_path))
    assert "Name | Age" in loaded.text
    assert "Alice | 30" in loaded.text
    assert loaded.metadata["rows"] == 3


def test_load_json_dict(tmp_path: Path) -> None:
    """JSON dict is flattened to key-value lines."""
    file_path = tmp_path / "sample.json"
    file_path.write_text(json.dumps({"name": "Alice", "age": 30}), encoding="utf-8")
    loaded = load_json(str(file_path))
    assert "name: Alice" in loaded.text
    assert "age: 30" in loaded.text


def test_load_json_list(tmp_path: Path) -> None:
    """JSON list is flattened to one line per item."""
    file_path = tmp_path / "sample.json"
    file_path.write_text(json.dumps(["a", "b", "c"]), encoding="utf-8")
    loaded = load_json(str(file_path))
    assert loaded.text == "a\nb\nc"


def test_load_html_strips_tags(tmp_path: Path) -> None:
    """HTML tags, scripts and styles are stripped."""
    file_path = tmp_path / "sample.html"
    html = (
        "<html><head><script>alert('x')</script><style>.x{}</style></head>"
        "<body><h1>Title</h1><p>Paragraph</p></body></html>"
    )
    file_path.write_text(html, encoding="utf-8")
    loaded = load_html(str(file_path))
    assert "<script>" not in loaded.text
    assert "<style>" not in loaded.text
    assert "<h1>" not in loaded.text
    assert "Title" in loaded.text
    assert "Paragraph" in loaded.text


def test_load_file_auto_detect(tmp_path: Path) -> None:
    """load_file dispatches to the correct loader by extension."""
    file_path = tmp_path / "auto.txt"
    file_path.write_text("dispatch test", encoding="utf-8")
    loaded = load_file(str(file_path))
    assert loaded.text == "dispatch test"
    assert loaded.file_type == "text"


def test_load_file_unsupported_extension(tmp_path: Path) -> None:
    """Unsupported extensions raise ValueError."""
    file_path = tmp_path / "unsupported.xyz"
    file_path.write_text("data", encoding="utf-8")
    with pytest.raises(ValueError, match="Unsupported file type"):
        load_file(str(file_path))


def test_load_file_missing_file(tmp_path: Path) -> None:
    """Missing files raise FileNotFoundError."""
    with pytest.raises(FileNotFoundError):
        load_file(str(tmp_path / "does_not_exist.txt"))
