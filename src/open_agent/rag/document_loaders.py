"""Document loaders for various file formats (PDF, DOCX, Markdown, TXT, CSV, JSON)."""
from __future__ import annotations

import csv
import json
from dataclasses import dataclass
from pathlib import Path
from typing import Any


@dataclass
class LoadedDocument:
    """A document loaded from a file, ready for indexing."""
    id: str
    text: str
    metadata: dict[str, Any]
    source: str
    file_type: str


def load_file(file_path: str) -> LoadedDocument:
    """Auto-detect file type and load using the appropriate loader.

    Supports: .txt, .md, .rst, .pdf, .docx, .csv, .json, .html, .htm
    Raises ValueError for unsupported file types.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = path.suffix.lower()
    loaders = {
        ".txt": load_text,
        ".md": load_markdown,
        ".rst": load_text,
        ".text": load_text,
        ".pdf": load_pdf,
        ".docx": load_docx,
        ".csv": load_csv,
        ".json": load_json,
        ".html": load_html,
        ".htm": load_html,
    }

    loader = loaders.get(ext)
    if loader is None:
        raise ValueError(f"Unsupported file type: {ext}")

    return loader(str(path))


def load_text(file_path: str) -> LoadedDocument:
    """Load a plain text file."""
    path = Path(file_path)
    text = path.read_text(encoding="utf-8", errors="replace")
    return LoadedDocument(
        id=path.stem,
        text=text,
        metadata={"source": file_path, "file_type": "text"},
        source=file_path,
        file_type="text",
    )


def load_markdown(file_path: str) -> LoadedDocument:
    """Load a Markdown file, stripping frontmatter if present."""
    path = Path(file_path)
    text = path.read_text(encoding="utf-8", errors="replace")
    # Strip YAML frontmatter
    if text.startswith("---"):
        parts = text.split("---", 2)
        if len(parts) >= 3:
            text = parts[2].strip()
    return LoadedDocument(
        id=path.stem,
        text=text,
        metadata={"source": file_path, "file_type": "markdown"},
        source=file_path,
        file_type="markdown",
    )


def load_pdf(file_path: str) -> LoadedDocument:
    """Load a PDF file using PyMuPDF (fitz).

    Preserves page boundaries and extracts tables when PyMuPDF supports it.
    Falls back to PyPDF2 if PyMuPDF is not installed.
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return _load_pdf_with_pypdf2(file_path)

    doc = fitz.open(file_path)
    path = Path(file_path)

    pages_text: list[str] = []
    tables: list[list[list[str]]] = []
    for page_num, page in enumerate(doc, start=1):
        page_text = page.get_text().strip()
        # Try to extract tables from the page.
        try:
            tab = page.find_tables()
            if tab and tab.tables:
                for table in tab.tables:
                    rows = table.extract()
                    if rows:
                        tables.append(rows)
                        # Append a textual representation of the table.
                        lines = [" | ".join(str(cell) for cell in row) for row in rows]
                        page_text += "\n\n[Table]\n" + "\n".join(lines)
        except Exception:
            # Table extraction may not be available in older PyMuPDF versions.
            pass

        if page_text:
            pages_text.append(f"--- Page {page_num} ---\n\n{page_text}")

    text = "\n\n".join(pages_text)
    metadata = {
        "source": file_path,
        "file_type": "pdf",
        "pages": doc.page_count,
        "title": doc.metadata.get("title") if doc.metadata else None,
        "author": doc.metadata.get("author") if doc.metadata else None,
        "tables": len(tables),
    }
    result = LoadedDocument(
        id=path.stem,
        text=text,
        metadata={k: v for k, v in metadata.items() if v is not None},
        source=file_path,
        file_type="pdf",
    )
    doc.close()
    return result


def _load_pdf_with_pypdf2(file_path: str) -> LoadedDocument:
    """Fallback PDF loader using PyPDF2."""
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        raise ImportError(
            "PDF loading requires PyMuPDF (pip install pymupdf) or PyPDF2 (pip install pypdf2)."
        )

    reader = PdfReader(file_path)
    pages_text = []
    for page_num, page in enumerate(reader.pages, start=1):
        page_text = (page.extract_text() or "").strip()
        if page_text:
            pages_text.append(f"--- Page {page_num} ---\n\n{page_text}")
    text = "\n\n".join(pages_text)
    path = Path(file_path)
    return LoadedDocument(
        id=path.stem,
        text=text,
        metadata={"source": file_path, "file_type": "pdf", "pages": len(reader.pages)},
        source=file_path,
        file_type="pdf",
    )


def load_docx(file_path: str) -> LoadedDocument:
    """Load a DOCX file using python-docx.

    Preserves headings, extracts tables, and records document properties.
    """
    try:
        import docx
    except ImportError:
        raise ImportError(
            "DOCX loading requires python-docx (pip install python-docx)."
        )

    doc = docx.Document(file_path)
    path = Path(file_path)

    # Extract paragraphs with heading markers.
    paragraphs: list[str] = []
    headings: list[dict[str, Any]] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_obj = para.style
        style = style_obj.name if style_obj is not None else "Normal"
        if style.startswith("Heading"):
            level = (
                int(style.replace("Heading ", ""))
                if style.replace("Heading ", "").isdigit()
                else 1
            )
            marker = "#" * level
            paragraphs.append(f"{marker} {text}")
            headings.append({"level": level, "text": text})
        else:
            paragraphs.append(text)

    # Extract tables as text blocks.
    tables: list[list[list[str]]] = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        if rows:
            tables.append(rows)
            lines = [" | ".join(cells) for cells in rows]
            paragraphs.append("[Table]\n" + "\n".join(lines))

    text = "\n\n".join(paragraphs)

    # Document core properties (title, author, etc.).
    props = doc.core_properties
    metadata: dict[str, Any] = {
        "source": file_path,
        "file_type": "docx",
        "paragraphs": len(doc.paragraphs),
        "headings": len(headings),
        "tables": len(tables),
    }
    if props.title:
        metadata["title"] = props.title
    if props.author:
        metadata["author"] = props.author
    if headings:
        metadata["heading_titles"] = [h["text"] for h in headings[:20]]

    return LoadedDocument(
        id=path.stem,
        text=text,
        metadata=metadata,
        source=file_path,
        file_type="docx",
    )


def load_csv(file_path: str) -> LoadedDocument:
    """Load a CSV file, converting rows to text."""
    path = Path(file_path)
    with open(file_path, newline="", encoding="utf-8", errors="replace") as f:
        reader = csv.reader(f)
        rows = list(reader)
    if not rows:
        text = ""
    else:
        header = rows[0]
        text_lines = [" | ".join(header)]
        text_lines.append("-" * len(text_lines[0]))
        for row in rows[1:]:
            text_lines.append(" | ".join(row))
        text = "\n".join(text_lines)
    return LoadedDocument(
        id=path.stem,
        text=text,
        metadata={"source": file_path, "file_type": "csv", "rows": len(rows)},
        source=file_path,
        file_type="csv",
    )


def load_json(file_path: str) -> LoadedDocument:
    """Load a JSON file, flattening to text."""
    path = Path(file_path)
    data = json.loads(path.read_text(encoding="utf-8", errors="replace"))
    if isinstance(data, dict):
        text = "\n".join(f"{k}: {v}" for k, v in data.items())
    elif isinstance(data, list):
        text = "\n".join(str(item) for item in data)
    else:
        text = str(data)
    return LoadedDocument(
        id=path.stem,
        text=text,
        metadata={"source": file_path, "file_type": "json"},
        source=file_path,
        file_type="json",
    )


def load_html(file_path: str) -> LoadedDocument:
    """Load an HTML file, stripping tags."""
    path = Path(file_path)
    text = path.read_text(encoding="utf-8", errors="replace")
    # Simple tag stripping
    import re
    # Remove scripts and styles
    text = re.sub(r"<(script|style)[^>]*>.*?</\1>", "", text, flags=re.DOTALL | re.IGNORECASE)
    # Remove tags
    text = re.sub(r"<[^>]+>", "", text)
    # Clean whitespace
    text = re.sub(r"\n\s*\n", "\n\n", text.strip())
    return LoadedDocument(
        id=path.stem,
        text=text,
        metadata={"source": file_path, "file_type": "html"},
        source=file_path,
        file_type="html",
    )


SUPPORTED_EXTENSIONS = {
    ".txt", ".md", ".rst", ".text", ".pdf", ".docx", ".csv", ".json", ".html", ".htm",
}
